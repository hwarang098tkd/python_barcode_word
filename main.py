from datetime import datetime
from random import randrange
import barcode
import pyodbc
import os
from barcode.writer import ImageWriter
from copy import deepcopy
from docx import Document
from docx.shared import Inches
import json

class Connection:

    def __init__(self):
        #get the json with login info
        with open('server_cred.json', 'r', encoding='utf-8') as f:
            credential_json = json.load(f)
        # initialize the objects
        self.username = credential_json["username"]
        self.password = credential_json["password"]
        self.server = credential_json["hostname"]
        self.database = credential_json["database_name"]
        # establishe a new connection
        try:
            self.cnxn = pyodbc.connect(
                'DRIVER={ODBC Driver 17 for SQL Server};SERVER=' + self.server +
                ';DATABASE=' + self.database + ';UID=' + self.username + ';PWD=' + self.password  +
                ';Trusted_Connection=no', timeout=20)
            print('Connection established')

        except Exception as e:
            print(str(e))

    # ####################################
    def get_old_barcodes(self):
        try:
            #
            cursor = self.cnxn.cursor()
            cursor.execute("SELECT BARCODE, BARCODE_1 FROM table_name")
            msg = cursor.fetchall()
        except Exception as e:
            msg = "Connection failed"
            print(str(e))
        return msg


def modify_list(list):
    # Modify the old list to new one , old list has inside lists with two items each
    new_list = []
    for barcode in list:
        new_list.append(barcode[0])
        new_list.append(barcode[1])
    return new_list


def barcode_images_creator(list):
    folder_name = date_time
    # check if the directory exists
    if not os.path.exists("barcodes_images/" + folder_name):
        os.makedirs("barcodes_images/" + folder_name)
        print("Created Directory : ", folder_name)
    else:
        print("Directory already existed : ", folder_name)
    # for each item in barcode list create a png file with a new barcode
    for item in list:
        barcode_new = barcode.get('ean13', str(item), writer=ImageWriter())
        barcode_new.save("barcodes_images/" + folder_name + "/" + str(item))
    print("Folder with barcodes images created: " + "barcodes_images/" + folder_name)

def create_word():
    # create a nwe word document
    def_doc = Document("words_documents/default_document.docx")
    # copy the content of the default document
    copy_the_content = deepcopy(def_doc)
    # create a new doc with the default copy
    copy_the_content.save("words_documents/" + date_time + ".docx")

def modify_word(list):
    # import the new word document
    new_doc = open("words_documents/" + date_time + ".docx", 'rb')
    document = Document(new_doc)
    y = 3
    x = 7
    # loop through all word's table cells and paste the barcode image
    for x1 in range(x):
        for y1 in range(y):
            document.tables[0].cell(x1, y1).add_paragraph().add_run().add_picture(
                "barcodes_images/" + date_time + "/" + str(list[0]) + ".png", width=Inches(2.7))
            # remove the inserted barcode image from barcode list
            list.pop(0)
    # save the dosument
    document.save("words_documents/" + date_time + ".docx")
    # and close it
    new_doc.close()

def main():
    # connect to database in order to collect existing barcodes
    conn = Connection()
    barcode_list = conn.get_old_barcodes()
    # call modify_list function
    old_barcodes = modify_list(barcode_list)
    word_barcodes = []
    counter = 0
    # generate barcodes randomly while checking if already exists
    while counter < 21:
        random_barc = randrange(1000000000000, 9999999999999, 2)
        if random_barc not in old_barcodes:
            word_barcodes.append(random_barc)
            counter += 1

    print(word_barcodes)
    # call the barcode_images_creator function in order ro create new barcodes
    barcode_images_creator(word_barcodes)
    # call the create_word function inn order to create a new doc with the default content
    create_word()
    # call the modify_word function to insert the barcode images
    modify_word(word_barcodes)

# Press the green button in the gutter to run the script.
if __name__ == '__main__':
    # create a folder name based on current datetime
    now = datetime.now()
    date_time = now.strftime("%m_%d_%Y_%H_%M_%S")
    main()
    print("Program terminated")

